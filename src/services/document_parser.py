import os
import tempfile
import zipfile
from typing import Dict, List, Any, Optional
import xml.etree.ElementTree as ET
from pathlib import Path
import re

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import markdown
except ImportError:
    markdown = None

class DocumentParser:
    """
    Service for parsing various document formats commonly used in healthcare software requirements.
    Supports PDF, Word documents, XML, Markdown, and plain text.
    """
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.xml', '.md', '.txt']
    
    def parse_document(self, file_path: str, file_content: bytes = None) -> Dict[str, Any]:
        """
        Parse a document and extract structured content.
        
        Args:
            file_path: Path to the document file
            file_content: Raw file content (optional, for uploaded files)
        
        Returns:
            Dictionary containing extracted content and metadata
        """
        if file_content:
            # Handle uploaded file content
            return self._parse_from_content(file_path, file_content)
        else:
            # Handle file path
            return self._parse_from_path(file_path)
    
    def _parse_from_path(self, file_path: str) -> Dict[str, Any]:
        """Parse document from file path."""
        if not os.path.exists(file_path):
            return {"error": "File not found", "content": "", "metadata": {}}
        
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_word(file_path)
            elif file_extension == '.xml':
                return self._parse_xml(file_path)
            elif file_extension == '.md':
                return self._parse_markdown(file_path)
            elif file_extension == '.txt':
                return self._parse_text(file_path)
            else:
                return {"error": f"Unsupported file format: {file_extension}", "content": "", "metadata": {}}
        except Exception as e:
            return {"error": f"Error parsing file: {str(e)}", "content": "", "metadata": {}}
    
    def _parse_from_content(self, filename: str, content: bytes) -> Dict[str, Any]:
        """Parse document from raw content."""
        file_extension = Path(filename).suffix.lower()
        
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name
            
            # Parse the temporary file
            result = self._parse_from_path(temp_path)
            
            # Clean up
            os.unlink(temp_path)
            
            return result
        except Exception as e:
            return {"error": f"Error processing uploaded file: {str(e)}", "content": "", "metadata": {}}
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF document."""
        if PyPDF2 is None:
            return {"error": "PyPDF2 not available for PDF parsing", "content": "", "metadata": {}}
        
        try:
            content = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'pages': len(pdf_reader.pages)
                    }
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        content += f"\\n--- Page {page_num + 1} ---\\n{page_text}\\n"
                    except Exception as e:
                        content += f"\\n--- Page {page_num + 1} (Error: {str(e)}) ---\\n"
            
            return {
                "content": content.strip(),
                "metadata": metadata,
                "format": "pdf"
            }
        except Exception as e:
            return {"error": f"Error parsing PDF: {str(e)}", "content": "", "metadata": {}}
    
    def _parse_word(self, file_path: str) -> Dict[str, Any]:
        """Parse Word document (.docx)."""
        if Document is None:
            return {"error": "python-docx not available for Word document parsing", "content": "", "metadata": {}}
        
        try:
            doc = Document(file_path)
            
            # Extract text content
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\\n"
            
            # Extract tables
            for table in doc.tables:
                content += "\\n--- Table ---\\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    content += row_text + "\\n"
                content += "--- End Table ---\\n"
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            return {
                "content": content.strip(),
                "metadata": metadata,
                "format": "docx"
            }
        except Exception as e:
            return {"error": f"Error parsing Word document: {str(e)}", "content": "", "metadata": {}}
    
    def _parse_xml(self, file_path: str) -> Dict[str, Any]:
        """Parse XML document."""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Extract text content recursively
            def extract_text(element, level=0):
                text = ""
                indent = "  " * level
                
                if element.text and element.text.strip():
                    text += f"{indent}{element.tag}: {element.text.strip()}\\n"
                elif element.tag:
                    text += f"{indent}{element.tag}:\\n"
                
                for child in element:
                    text += extract_text(child, level + 1)
                
                return text
            
            content = extract_text(root)
            
            # Extract metadata
            metadata = {
                'root_tag': root.tag,
                'namespace': root.tag.split('}')[0][1:] if '}' in root.tag else '',
                'attributes': dict(root.attrib),
                'elements_count': len(list(root.iter()))
            }
            
            return {
                "content": content.strip(),
                "metadata": metadata,
                "format": "xml"
            }
        except Exception as e:
            return {"error": f"Error parsing XML: {str(e)}", "content": "", "metadata": {}}
    
    def _parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """Parse Markdown document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Extract metadata (if markdown has front matter)
            metadata = {}
            if content.startswith('---'):
                parts = content.split('---', 2)
                if len(parts) >= 3:
                    # Parse YAML front matter (basic parsing)
                    front_matter = parts[1].strip()
                    for line in front_matter.split('\\n'):
                        if ':' in line:
                            key, value = line.split(':', 1)
                            metadata[key.strip()] = value.strip()
                    content = parts[2].strip()
            
            # Extract headers for structure
            headers = re.findall(r'^(#{1,6})\\s+(.+)$', content, re.MULTILINE)
            metadata['headers'] = [{'level': len(h[0]), 'text': h[1]} for h in headers]
            metadata['format'] = 'markdown'
            
            return {
                "content": content,
                "metadata": metadata,
                "format": "markdown"
            }
        except Exception as e:
            return {"error": f"Error parsing Markdown: {str(e)}", "content": "", "metadata": {}}
    
    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """Parse plain text document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Basic metadata
            metadata = {
                'lines': len(content.split('\\n')),
                'characters': len(content),
                'words': len(content.split()),
                'format': 'text'
            }
            
            return {
                "content": content,
                "metadata": metadata,
                "format": "text"
            }
        except Exception as e:
            return {"error": f"Error parsing text file: {str(e)}", "content": "", "metadata": {}}
    
    def extract_requirements_sections(self, content: str) -> List[Dict[str, Any]]:
        """
        Extract potential requirements sections from document content.
        """
        sections = []
        
        # Common requirement section patterns
        patterns = [
            r'(?i)(?:^|\n)(\d+\.?\d*\.?\d*)\s+(.*?)(?=\n\d+\.|\n[A-Z][A-Z\s]+:|\Z)',
            r'(?i)(?:^|\n)(REQ-?\d+)\s+(.*?)(?=\nREQ-?\d+|\n[A-Z][A-Z\s]+:|\Z)',
            r'(?i)(?:^|\n)([A-Z]{2,}-\d+)\s+(.*?)(?=\n[A-Z]{2,}-\d+|\n[A-Z][A-Z\s]+:|\Z)',
            r'(?i)(?:^|\n)(requirement\s+\d+)\s*:?\s*(.*?)(?=\nrequirement\s+\d+|\n[A-Z][A-Z\s]+:|\Z)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, content, re.DOTALL | re.MULTILINE)
            for match in matches:
                if len(match) == 2:
                    req_id, req_text = match
                    sections.append({
                        'id': req_id.strip(),
                        'text': req_text.strip()[:1000],  # Limit text length
                        'type': 'requirement'
                    })
        
        # If no structured requirements found, split by paragraphs
        if not sections:
            paragraphs = [p.strip() for p in content.split('\\n\\n') if p.strip()]
            for i, paragraph in enumerate(paragraphs[:20]):  # Limit to first 20 paragraphs
                if len(paragraph) > 50:  # Only consider substantial paragraphs
                    sections.append({
                        'id': f'PARA-{i+1:03d}',
                        'text': paragraph,
                        'type': 'paragraph'
                    })
        
        return sections
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if the file format is supported."""
        file_extension = Path(filename).suffix.lower()
        return file_extension in self.supported_formats
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return self.supported_formats.copy()

